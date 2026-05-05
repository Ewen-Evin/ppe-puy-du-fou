#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur de documentation PPE Puy du Fou V1
Crée un DOCX complet + upload vers Google Docs (Drive API)

Usage:
    python create_documentation.py              # génère DOCX local seulement
    python create_documentation.py --upload     # génère + upload vers Google Drive
"""

import os, sys, json, io
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Chemins ─────────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MCD_IMAGE   = os.path.join(SCRIPT_DIR, '..', 'db', 'mcd.jpg')
OUTPUT_FILE = os.path.join(SCRIPT_DIR, 'PPE_PuyDuFou_Documentation.docx')

# ─── Palette couleurs ─────────────────────────────────────────────────────────
BLUE       = RGBColor(0x1a, 0x73, 0xe8)
GREEN      = RGBColor(0x0f, 0x9d, 0x58)
DARK_GRAY  = RGBColor(0x3c, 0x40, 0x43)
LIGHT_GRAY = RGBColor(0x80, 0x86, 0x8d)
WHITE      = RGBColor(0xff, 0xff, 0xff)
CODE_BG    = RGBColor(0xf1, 0xf3, 0xf4)
HEADER_BG  = RGBColor(0x1a, 0x73, 0xe8)

# ─── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_col_widths(table, widths):
    """Set column widths (in Cm) for a table."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = Cm(widths[i])


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = BLUE
        run.font.size = Pt(18)
    elif level == 2:
        run.font.color.rgb = GREEN
        run.font.size = Pt(14)
    elif level == 3:
        run.font.color.rgb = DARK_GRAY
        run.font.size = Pt(12)
    return p


def add_para(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold  = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_code(doc, text):
    """Add a grey code block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_GRAY
    # Background color via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F1F3F4')
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], '1A73E8')
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            row[i].paragraphs[0].runs[0].font.size = Pt(9.5)

    if col_widths:
        set_col_widths(table, col_widths)

    doc.add_paragraph()
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(level * 0.5)
    return p


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE DE COUVERTURE
# ═══════════════════════════════════════════════════════════════════════════════

def build_cover(doc):
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PPE PUY DU FOU')
    run.font.size = Pt(38)
    run.font.bold = True
    run.font.color.rgb = BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Documentation Complète du Projet')
    run2.font.size = Pt(20)
    run2.font.color.rgb = DARK_GRAY

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('API REST · Back-office Web · Application Android')
    run3.font.size = Pt(14)
    run3.font.color.rgb = LIGHT_GRAY
    run3.font.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run('Version 1.0  —  Mai 2026')
    run4.font.size = Pt(12)
    run4.font.color.rgb = LIGHT_GRAY

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# PRÉSENTATION DU PROJET
# ═══════════════════════════════════════════════════════════════════════════════

def build_presentation(doc):
    add_heading(doc, '1. Présentation du Projet', 1)

    add_heading(doc, '1.1 Contexte', 2)
    add_para(doc, (
        'Le projet PPE Puy du Fou est une application de planification de visite pour le parc '
        'Puy du Fou (Vendée, France). Il s\'agit de la version 1 (V1), refonte complète de la V0, '
        'qui sépare les responsabilités en trois composantes indépendantes communicant via une '
        'API REST centralisée.'
    ))

    add_heading(doc, '1.2 Objectifs', 2)
    add_bullet(doc, 'Permettre aux visiteurs de planifier leur journée au parc via une application mobile Android.')
    add_bullet(doc, 'Calculer automatiquement des parcours optimisés selon les spectacles choisis, la vitesse de marche et les horaires.')
    add_bullet(doc, 'Offrir aux gestionnaires du parc un back-office web pour administrer spectacles, séances, lieux et jours d\'ouverture.')
    add_bullet(doc, 'Centraliser toutes les données et la logique métier dans une API REST sécurisée par JWT.')

    doc.add_paragraph()

    add_heading(doc, '1.3 Architecture Globale', 2)
    add_para(doc, 'Le système est composé de trois couches :')

    add_table(doc,
        ['Composante', 'Technologie', 'Rôle', 'URL'],
        [
            ['API REST', 'PHP 8 MVC', 'Logique métier + BDD', 'localhost/ppe-puy-du-fou/V1/api/public'],
            ['Back-office', 'PHP 8 MVC + cURL', 'Interface de gestion', 'localhost/ppe-puy-du-fou/V1/backoffice/public'],
            ['App Android', 'Java (Android SDK 24+)', 'Interface visiteur', 'Application APK'],
        ],
        col_widths=[3, 3.5, 4, 6]
    )

    add_para(doc, (
        'L\'API est le seul composant qui accède directement à la base de données MySQL. '
        'Le back-office et l\'application Android consomment l\'API via HTTP (JSON). '
        'L\'authentification est gérée par tokens JWT (HS256, durée 24h).'
    ))

    add_heading(doc, '1.4 Technologies', 2)
    add_table(doc,
        ['Technologie', 'Version', 'Usage'],
        [
            ['PHP', '8.x', 'Backend API + Back-office'],
            ['MySQL', '8', 'Base de données relationnelle'],
            ['PDO', 'Built-in PHP', 'Requêtes SQL préparées (anti-injection)'],
            ['JWT (HS256)', 'Maison', 'Authentification sans état'],
            ['Java', 'JDK 11+', 'Application Android'],
            ['Android SDK', 'API 24 (Android 8+)', 'Plateforme mobile'],
            ['OSMDroid', 'Latest', 'Cartographie OpenStreetMap'],
            ['cURL', 'PHP extension', 'Appels HTTP back-office → API'],
            ['Laragon', 'Latest', 'Stack LAMP Windows (dev local)'],
        ],
        col_widths=[4, 3, 9]
    )

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# MCD
# ═══════════════════════════════════════════════════════════════════════════════

def build_mcd(doc):
    add_heading(doc, '2. Modèle Conceptuel de Données (MCD)', 1)

    add_para(doc, (
        'Le MCD représente les entités métier du système et leurs relations. '
        'La base de données contient 10 tables pour gérer les visiteurs, les spectacles, '
        'les séances, les visites et les parcours calculés.'
    ))

    # Image
    if os.path.exists(MCD_IMAGE):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(MCD_IMAGE, width=Inches(6.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = cap.add_run('Figure 1 — Modèle Conceptuel de Données (Merise)')
        run_cap.font.size = Pt(10)
        run_cap.font.italic = True
        run_cap.font.color.rgb = LIGHT_GRAY
        doc.add_paragraph()
    else:
        add_para(doc, '[Image MCD non trouvée — placer db/mcd.jpg]', italic=True, color=LIGHT_GRAY)

    add_heading(doc, '2.1 Entités et Attributs', 2)

    add_table(doc,
        ['Entité (Table)', 'Clé Primaire', 'Attributs principaux', 'Description'],
        [
            ['utilisateur', 'id_utilisateur', 'email, nom, prenom, mot_de_passe (bcrypt), vitesse_marche, type_profil', 'Comptes visiteurs et gestionnaires'],
            ['spectacle', 'id_spectacle', 'libelle, description, duree_spectacle, duree_attente, id_lieu', '11 spectacles du parc'],
            ['lieu', 'id_lieu', 'nom_lieu, type_lieu (ENUM), coordonnees_gps (lat,lng)', '40 lieux : spectacles, zones, restaurants, hôtels'],
            ['seance', 'id_seance', 'date_seance, heure_debut, heure_fin, id_spectacle', '36 séances réparties sur 4 jours'],
            ['jours', 'id_jours (DATE)', 'heure_ouverture, heure_fermeture', 'Jours d\'ouverture du parc'],
            ['visite', 'id_visite', 'nom_visite, date_visite, vitesse_marche, id_utilisateur', 'Planification de visite d\'un utilisateur'],
            ['choisir', '(id_spectacle, id_visite)', '—', 'Table de jonction : spectacles choisis pour une visite'],
            ['parcours', 'id_parcours', 'duree, est_complet, temps_attente, est_favori, id_visite', 'Résultat du calcul d\'itinéraire'],
            ['etape', 'id_etape', 'ordre, heure_arrivee, id_parcours, id_seance', 'Étapes ordonnées d\'un parcours'],
            ['distance', '(id_lieu, id_lieu_1)', 'distance_metres', 'Graphe des distances entre lieux (144 arêtes)'],
        ],
        col_widths=[3, 3.5, 7, 4]
    )

    add_heading(doc, '2.2 Relations', 2)

    add_table(doc,
        ['Relation', 'Cardinalité', 'Description'],
        [
            ['utilisateur → visite', '1,N', 'Un utilisateur possède plusieurs visites'],
            ['visite → choisir → spectacle', 'N,N', 'Une visite sélectionne plusieurs spectacles'],
            ['spectacle → seance', '1,N', 'Un spectacle a plusieurs séances (sur plusieurs jours)'],
            ['lieu → spectacle', '1,N', 'Un lieu accueille un spectacle'],
            ['lieu ↔ distance ↔ lieu', 'N,N', 'Graphe bidirectionnel de distances entre lieux'],
            ['visite → parcours', '1,N', 'Une visite génère plusieurs parcours alternatifs'],
            ['parcours → etape → seance', 'N,N', 'Un parcours est composé d\'étapes liées à des séances'],
            ['jours → seance', '1,N', 'Un jour d\'ouverture contient plusieurs séances'],
        ],
        col_widths=[5, 3, 9]
    )

    add_heading(doc, '2.3 Contraintes d\'intégrité', 2)
    add_bullet(doc, 'Email utilisateur : UNIQUE (contrainte BDD + validation API).')
    add_bullet(doc, 'Clés étrangères avec ON DELETE CASCADE sur : visite, spectacle, seance, parcours, choisir, distance.')
    add_bullet(doc, 'type_profil : ENUM(\'visiteur\', \'gestionnaire\') avec valeur par défaut \'visiteur\'.')
    add_bullet(doc, 'Encodage : UTF-8MB4 sur toutes les tables (support emojis et accents).')
    add_bullet(doc, 'Index sur : seance.date_seance, spectacle.id_lieu, utilisateur.email.')

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# DOC TECHNIQUE — API
# ═══════════════════════════════════════════════════════════════════════════════

def build_api(doc):
    add_heading(doc, '3. Documentation Technique', 1)
    add_heading(doc, '3.1 API REST', 2)

    add_para(doc, (
        'L\'API REST est le cœur du système. Elle centralise toute la logique métier, '
        'les accès à la base de données et le calcul d\'itinéraires. Elle expose 46 endpoints '
        'répartis en 3 groupes selon le niveau d\'authentification requis.'
    ))

    add_heading(doc, '3.1.1 Architecture & Structure', 3)

    add_para(doc, 'Structure du projet API :')
    add_code(doc, 'api/')
    add_code(doc, '├── public/index.php          # Point d\'entrée unique (front-controller)')
    add_code(doc, '├── config/config.php          # Configuration BDD + JWT')
    add_code(doc, '└── app/')
    add_code(doc, '    ├── Controllers/           # 6 contrôleurs métier')
    add_code(doc, '    ├── Core/                  # Router, Database (PDO), Jwt, Controller, Middleware')
    add_code(doc, '    ├── Middlewares/           # AuthMiddleware (JWT + rôles)')
    add_code(doc, '    ├── Models/                # 8 modèles (accès BDD)')
    add_code(doc, '    ├── Services/              # ParcoursService (algorithme calcul)')
    add_code(doc, '    └── Routes/routes.php      # 46 routes')
    doc.add_paragraph()

    add_para(doc, 'Le Router maison analyse la méthode HTTP et le path, extrait les paramètres '
        'dynamiques ({id}, {date}), exécute la chaîne de middlewares puis appelle le contrôleur cible.')

    add_heading(doc, '3.1.2 Authentification JWT', 3)

    add_para(doc, 'L\'API utilise des tokens JWT (JSON Web Token) signés en HS256 sans dépendance externe.')

    add_table(doc,
        ['Propriété', 'Valeur'],
        [
            ['Algorithme', 'HS256'],
            ['Durée de validité', '86400 secondes (24 heures)'],
            ['Payload', 'sub (id utilisateur), role, iat, exp, iss'],
            ['Header', 'Authorization: Bearer <token>'],
            ['Stockage Android', 'SharedPreferences'],
            ['Stockage Back-office', '$_SESSION[\'token\'] (PHP session côté serveur)'],
        ],
        col_widths=[5, 12]
    )

    add_para(doc, 'AuthMiddleware vérifie la signature et l\'expiration du token, puis contrôle le rôle '
        '(visiteur ou gestionnaire) selon la route. Renvoie 401 si token absent/invalide, 403 si rôle insuffisant.')

    add_heading(doc, '3.1.3 Endpoints — Routes Publiques', 3)
    add_para(doc, 'Ces routes ne nécessitent aucun token. Accessibles depuis l\'application Android et le back-office.')

    add_table(doc,
        ['Méthode', 'Route', 'Description'],
        [
            ['GET',  '/api/health',                    'Vérification que l\'API est en ligne'],
            ['POST', '/api/auth/register',             'Inscription visiteur (email, mot_de_passe, nom, prenom, vitesse_marche)'],
            ['POST', '/api/auth/login',                'Connexion — retourne JWT + données utilisateur'],
            ['GET',  '/api/spectacles',                'Liste tous les spectacles avec le nom et les GPS du lieu associé'],
            ['GET',  '/api/spectacles/{id}',           'Détail d\'un spectacle'],
            ['GET',  '/api/lieux',                     'Liste les 40 lieux avec coordonnées GPS'],
            ['GET',  '/api/lieux/{id}',                'Détail d\'un lieu'],
            ['GET',  '/api/distances',                 'Graphe complet de distances (144 arêtes)'],
            ['GET',  '/api/seances',                   'Liste des séances (filtre optionnel ?date=YYYY-MM-DD)'],
            ['GET',  '/api/spectacles/{id}/seances',   'Séances d\'un spectacle'],
            ['GET',  '/api/jours',                     'Jours d\'ouverture du parc'],
            ['GET',  '/api/jours/{date}',              'Horaires d\'un jour spécifique'],
        ],
        col_widths=[2.5, 6, 8.5]
    )

    add_heading(doc, '3.1.4 Endpoints — Routes Visiteur (JWT requis)', 3)

    add_table(doc,
        ['Méthode', 'Route', 'Description'],
        [
            ['GET',    '/api/auth/me',                'Profil de l\'utilisateur connecté'],
            ['PUT',    '/api/auth/vitesse',           'Mise à jour de la vitesse de marche'],
            ['POST',   '/api/parcours/preview',       'Calcule des parcours SANS sauvegarder en BDD'],
            ['POST',   '/api/visites',                'Crée une visite et calcule + enregistre les parcours'],
            ['GET',    '/api/visites',                'Historique des visites de l\'utilisateur'],
            ['GET',    '/api/visites/{id}/parcours',  'Parcours d\'une visite (recalcule si absent)'],
            ['GET',    '/api/visites/{id}/carte',     'Points GPS du parcours favori (pour la carte OSM)'],
            ['PUT',    '/api/visites/{id}/favori',    'Marque un parcours comme favori'],
            ['DELETE', '/api/visites/{id}',           'Supprime une visite et ses parcours associés'],
        ],
        col_widths=[2.5, 6, 8.5]
    )

    add_heading(doc, '3.1.5 Endpoints — Routes Gestionnaire (JWT gestionnaire requis)', 3)

    add_table(doc,
        ['Méthode', 'Route', 'Description'],
        [
            ['POST',   '/api/admin/spectacles',          'Créer un spectacle'],
            ['PUT',    '/api/admin/spectacles/{id}',     'Modifier un spectacle'],
            ['DELETE', '/api/admin/spectacles/{id}',     'Supprimer un spectacle'],
            ['POST',   '/api/admin/lieux',               'Créer un lieu'],
            ['PUT',    '/api/admin/lieux/{id}',          'Modifier un lieu'],
            ['DELETE', '/api/admin/lieux/{id}',          'Supprimer un lieu'],
            ['POST',   '/api/admin/distances',           'Ajouter une arête dans le graphe (bidirectionnel)'],
            ['DELETE', '/api/admin/distances/{a}/{b}',   'Supprimer une arête du graphe'],
            ['POST',   '/api/admin/seances',             'Créer une séance'],
            ['PUT',    '/api/admin/seances/{id}',        'Modifier une séance'],
            ['DELETE', '/api/admin/seances/{id}',        'Supprimer une séance'],
            ['POST',   '/api/admin/jours',               'Créer ou mettre à jour un jour d\'ouverture (UPSERT)'],
            ['DELETE', '/api/admin/jours/{date}',        'Supprimer un jour d\'ouverture'],
        ],
        col_widths=[2.5, 6.5, 8]
    )

    add_heading(doc, '3.1.6 Exemples de Requêtes', 3)

    add_para(doc, 'Inscription visiteur :', bold=True)
    add_code(doc, 'POST /api/auth/register')
    add_code(doc, 'Content-Type: application/json')
    add_code(doc, '')
    add_code(doc, '{')
    add_code(doc, '  "email":          "visiteur@test.com",')
    add_code(doc, '  "mot_de_passe":   "MotDePasse123",')
    add_code(doc, '  "nom":            "Dupont",')
    add_code(doc, '  "prenom":         "Jean",')
    add_code(doc, '  "vitesse_marche": 4.5')
    add_code(doc, '}')
    add_code(doc, '→ HTTP 201  { "id_utilisateur": 4 }')
    doc.add_paragraph()

    add_para(doc, 'Calcul de parcours (preview) :', bold=True)
    add_code(doc, 'POST /api/parcours/preview')
    add_code(doc, 'Authorization: Bearer <token>')
    add_code(doc, '')
    add_code(doc, '{')
    add_code(doc, '  "date_visite":    "2026-04-11",')
    add_code(doc, '  "vitesse_marche": 4.5,')
    add_code(doc, '  "spectacles":     [1, 6, 9]')
    add_code(doc, '}')
    add_code(doc, '→ HTTP 200  { "parcours": [ { "etapes": [...], "complet": true,')
    add_code(doc, '                              "duree_totale_min": 190, "attente_min": 110 }, ... ] }')
    doc.add_paragraph()

    add_heading(doc, '3.1.7 Service de Calcul de Parcours (ParcoursService)', 3)

    add_para(doc, (
        'C\'est le cœur algorithmique du projet. Le service calcule tous les ordres de visite '
        'possibles et retourne les 10 meilleurs itinéraires, triés par complétude puis par temps '
        'd\'attente minimal.'
    ))

    add_para(doc, 'Étape 1 — Floyd-Warshall (construction du graphe des distances) :', bold=True)
    add_para(doc, (
        'L\'algorithme construit une matrice distance[A][B] à partir des 144 arêtes de la table '
        'distance. Il applique ensuite la transitivité de Floyd-Warshall pour calculer le plus '
        'court chemin entre tous les couples de lieux. Complexité : O(n³) sur ~40 nœuds.'
    ))

    add_para(doc, 'Étape 2 — Backtracking récursif :', bold=True)
    add_para(doc, (
        'Pour chaque permutation de spectacles choisis, l\'algorithme vérifie la faisabilité '
        'séance par séance : '
        '(1) le spectacle n\'est pas encore visité, '
        '(2) la séance est dans les horaires du parc, '
        '(3) le temps de trajet depuis la séance précédente (calculé via Floyd) laisse assez '
        'de marge avant l\'heure de début de la séance. '
        'Si toutes les conditions sont remplies, la séance est ajoutée et la récursion continue.'
    ))

    add_para(doc, 'Étape 3 — Fallback glouton (si aucun parcours complet) :', bold=True)
    add_para(doc, (
        'Si le backtracking ne trouve aucun parcours intégrant tous les spectacles, un algorithme '
        'glouton sélectionne le maximum de spectacles compatibles. Le parcours est alors marqué '
        'est_complet = false.'
    ))

    add_para(doc, 'Étape 4 — Finalisation et tri :', bold=True)
    add_para(doc, (
        'Chaque parcours candidat est enrichi des durées (totale, attente), puis les résultats '
        'sont triés : parcours complets en premier, puis par temps d\'attente croissant. '
        'Les 10 meilleurs parcours sont retournés.'
    ))

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# DOC TECHNIQUE — BACKOFFICE
# ═══════════════════════════════════════════════════════════════════════════════

def build_backoffice(doc):
    add_heading(doc, '3.2 Back-office Web', 2)

    add_para(doc, (
        'Le back-office est une interface web PHP accessible aux gestionnaires du parc. '
        'Il consomme exclusivement l\'API REST via cURL. Aucun accès direct à la base de données.'
    ))

    add_heading(doc, '3.2.1 Architecture & Structure', 3)

    add_code(doc, 'backoffice/')
    add_code(doc, '├── public/index.php         # Front-controller (session PHP + CSRF)')
    add_code(doc, '├── config/config.php         # URL de l\'API')
    add_code(doc, '└── app/')
    add_code(doc, '    ├── Controllers/          # 7 contrôleurs')
    add_code(doc, '    ├── Core/                 # Router + ApiClient (cURL) + Controller')
    add_code(doc, '    ├── Routes/routes.php     # 31 routes')
    add_code(doc, '    └── Views/                # 11 templates PHP')
    doc.add_paragraph()

    add_heading(doc, '3.2.2 Authentification', 3)
    add_para(doc, (
        'La connexion utilise l\'endpoint POST /api/auth/login. Si le type_profil retourné est '
        '\'gestionnaire\', le token JWT est stocké dans $_SESSION[\'token\']. '
        'Toute page protégée vérifie l\'existence de cette session. '
        'La déconnexion détruit la session complète.'
    ))

    add_heading(doc, '3.2.3 Modules Fonctionnels', 3)

    add_table(doc,
        ['Module', 'Contrôleur', 'Fonctionnalités', 'Routes'],
        [
            ['Tableau de bord', 'DashboardController', 'Compteurs : nb spectacles, lieux, jours, séances', '1'],
            ['Spectacles', 'SpectaclesController', 'Liste, créer, éditer, supprimer un spectacle', '5'],
            ['Lieux', 'LieuxController', 'Liste, créer, éditer, supprimer un lieu (avec GPS)', '5'],
            ['Séances', 'SeancesController', 'Liste (filtre par date), créer, éditer, supprimer', '5'],
            ['Jours', 'JoursController', 'Afficher, créer/modifier (UPSERT), supprimer un jour', '4'],
            ['Distances', 'DistancesController', 'Afficher graphe, ajouter arête, supprimer arête', '3'],
            ['Auth', 'AuthController', 'Formulaire login, traitement login, logout', '3'],
        ],
        col_widths=[3, 4, 7.5, 1.5]
    )

    add_heading(doc, '3.2.4 Sécurité', 3)
    add_bullet(doc, 'Session PHP : démarrée à chaque requête, vérifiée sur toutes les routes protégées.')
    add_bullet(doc, 'Tokens CSRF : générés avec bin2hex(random_bytes(32)), vérifiés sur tous les POST.')
    add_bullet(doc, 'JWT : stocké côté serveur dans $_SESSION (non exposé au navigateur).')
    add_bullet(doc, 'Redirection automatique vers /login si session absente.')
    add_bullet(doc, 'Validation des codes HTTP retournés par l\'API avant affichage.')

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# DOC TECHNIQUE — ANDROID
# ═══════════════════════════════════════════════════════════════════════════════

def build_android(doc):
    add_heading(doc, '3.3 Application Android', 2)

    add_para(doc, (
        'L\'application Android est développée en Java natif pour Android SDK API 24 (Android 8.0+). '
        'Elle communique avec l\'API via HttpURLConnection sans bibliothèque tierce. '
        'La navigation repose sur 4 onglets Material 3 (BottomNavigationView).'
    ))

    add_heading(doc, '3.3.1 Activités (11 écrans)', 3)

    add_table(doc,
        ['Activité', 'Rôle', 'Endpoints appelés'],
        [
            ['LoginActivity', 'Authentification visiteur', 'POST /auth/login'],
            ['RegisterActivity', 'Création de compte', 'POST /auth/register'],
            ['MainActivity', 'Accueil : statut parc, prochaines séances, visites récentes', 'GET /jours, /seances?date, /visites'],
            ['SpectaclesActivity', 'Sélection des spectacles (date + checkbox)', 'GET /seances?date, POST /parcours/preview'],
            ['ParcoursActivity', 'Affichage et sélection d\'un parcours (preview ou historique)', 'GET /visites/{id}/parcours, PUT /visites/{id}/favori, POST /visites'],
            ['CatalogueActivity', 'Catalogue complet avec filtres (type de lieu)', 'GET /spectacles, /lieux'],
            ['SpectacleDetailActivity', 'Détail d\'un spectacle et ses séances futures', 'GET /spectacles/{id}/seances'],
            ['HistoriqueActivity', 'Liste des visites passées', 'GET /visites'],
            ['ProfileActivity', 'Profil utilisateur + slider vitesse de marche', 'PUT /auth/vitesse'],
            ['CarteActivity', 'Carte OSMDroid avec marqueurs GPS du parcours favori', 'GET /visites/{id}/carte'],
        ],
        col_widths=[4.5, 5.5, 7]
    )

    add_heading(doc, '3.3.2 Couche API — ApiClient', 3)
    add_para(doc, (
        'ApiClient.java centralise tous les appels HTTP. Il utilise un ExecutorService à 4 threads '
        'pour les requêtes réseau, puis publie les résultats sur le MainLooper (thread UI) via callback.'
    ))
    add_bullet(doc, 'Méthodes disponibles : get(), post(), put(), delete().')
    add_bullet(doc, 'Timeout : 10 secondes (connexion), 15 secondes (lecture).')
    add_bullet(doc, 'Token JWT injecté automatiquement depuis SharedPreferences (Authorization: Bearer).')
    add_bullet(doc, 'HTTP 401 : vide la session et redirige vers LoginActivity automatiquement.')
    add_bullet(doc, 'ApiResponse retourne le statut HTTP et le body JSON brut.')

    add_heading(doc, '3.3.3 Navigation', 3)
    add_para(doc, 'La navigation principale utilise BottomNavigationView avec 4 onglets :')
    add_table(doc,
        ['Onglet', 'Activité', 'Icône'],
        [
            ['Accueil', 'MainActivity', 'Maison'],
            ['Catalogue', 'CatalogueActivity', 'Liste/Recherche'],
            ['Historique', 'HistoriqueActivity', 'Horloge'],
            ['Profil', 'ProfileActivity', 'Personne'],
        ],
        col_widths=[4, 5, 5]
    )
    add_para(doc, 'NavHelper.java synchronise l\'état sélectionné du BottomNavigationView '
        'entre les 4 activités pour une navigation cohérente.')

    add_heading(doc, '3.3.4 Stockage Local', 3)
    add_table(doc,
        ['Classe', 'Mécanisme', 'Données stockées'],
        [
            ['Session', 'SharedPreferences', 'token JWT, id, email, nom, prenom, vitesse_marche, type_profil'],
            ['RecentlyViewed', 'SharedPreferences (JSON array)', 'Derniers 5 spectacles consultés (id, libelle, description, durée)'],
        ],
        col_widths=[3, 4.5, 9.5]
    )

    add_heading(doc, '3.3.5 Configuration Requise', 3)
    add_bullet(doc, 'Android 8.0 minimum (API 24).')
    add_bullet(doc, 'Permissions : INTERNET, ACCESS_NETWORK_STATE.')
    add_bullet(doc, 'Cleartext traffic activé (HTTP localhost en développement).')
    add_bullet(doc, 'Connexion réseau requise (Wi-Fi ou données mobiles sur le même réseau que le serveur).')

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENTATION UTILISATEUR — VISITEUR
# ═══════════════════════════════════════════════════════════════════════════════

def build_user_visitor(doc):
    add_heading(doc, '4. Documentation Utilisateur', 1)
    add_heading(doc, '4.1 Guide Visiteur — Application Android', 2)

    add_para(doc, (
        'Ce guide s\'adresse aux visiteurs du parc Puy du Fou qui utilisent l\'application Android '
        'pour planifier et organiser leur journée.'
    ))

    # --- Inscription
    add_heading(doc, '4.1.1 Inscription', 3)
    add_para(doc, 'Première utilisation : créer un compte depuis l\'écran de connexion.')
    add_table(doc,
        ['Étape', 'Action', 'Détail'],
        [
            ['1', 'Ouvrir l\'application', 'L\'écran de connexion s\'affiche automatiquement.'],
            ['2', 'Appuyer sur « Créer un compte »', 'Redirection vers l\'écran d\'inscription.'],
            ['3', 'Renseigner les informations', 'Prénom, Nom, Email, Mot de passe, Vitesse de marche.'],
            ['4', 'Valider', 'Appuyer sur « S\'inscrire ». Le compte est créé et la connexion automatique.'],
        ],
        col_widths=[1.5, 4.5, 11]
    )
    add_para(doc, 'Vitesse de marche : valeur par défaut 4,0 km/h. Peut être ajustée dans le profil.', italic=True)

    # --- Connexion
    add_heading(doc, '4.1.2 Connexion', 3)
    add_table(doc,
        ['Étape', 'Action'],
        [
            ['1', 'Saisir l\'adresse email et le mot de passe.'],
            ['2', 'Appuyer sur « Se connecter ».'],
            ['3', 'L\'application redirige vers l\'accueil si les identifiants sont corrects.'],
        ],
        col_widths=[1.5, 15.5]
    )

    # --- Accueil
    add_heading(doc, '4.1.3 Écran d\'Accueil', 3)
    add_para(doc, 'L\'accueil présente :')
    add_bullet(doc, 'Statut du parc : ouvert ou fermé selon le jour courant.')
    add_bullet(doc, 'Prochaines séances disponibles pour la date du jour.')
    add_bullet(doc, 'Derniers spectacles consultés (accès rapide au catalogue).')
    add_bullet(doc, 'Bouton « Planifier une visite » pour démarrer la planification.')

    # --- Catalogue
    add_heading(doc, '4.1.4 Catalogue des Spectacles', 3)
    add_para(doc, 'Accessible via l\'onglet « Catalogue » en bas de l\'écran.')
    add_bullet(doc, 'Affiche la liste de tous les spectacles du parc avec durée et lieu.')
    add_bullet(doc, 'Filtres disponibles : Tous, Spectacle, Zone, Restaurant, Hôtel.')
    add_bullet(doc, 'Appuyer sur un spectacle pour voir son détail (description, durée, prochaines séances).')

    # --- Planifier
    add_heading(doc, '4.1.5 Planifier une Visite', 3)
    add_para(doc, 'La planification se fait en plusieurs étapes :')
    add_table(doc,
        ['Étape', 'Action', 'Détail'],
        [
            ['1', 'Choisir la date de visite', 'Sélectionner parmi les jours d\'ouverture disponibles.'],
            ['2', 'Sélectionner les spectacles', 'Cocher les spectacles souhaités dans la liste.'],
            ['3', 'Appuyer sur « Calculer »', 'L\'application envoie la sélection à l\'API.'],
            ['4', 'Consulter les parcours proposés', 'Jusqu\'à 10 itinéraires sont affichés, triés par qualité.'],
            ['5', 'Choisir un itinéraire', 'Appuyer sur « Choisir cet itinéraire » sur le parcours souhaité.'],
            ['6', 'Nommer et enregistrer', 'Saisir un nom de visite et confirmer la sauvegarde.'],
        ],
        col_widths=[1.5, 4.5, 11]
    )

    add_para(doc, 'Les parcours sont calculés en tenant compte de :', bold=True)
    add_bullet(doc, 'Les horaires de chaque séance des spectacles sélectionnés.')
    add_bullet(doc, 'La distance entre les lieux et votre vitesse de marche personnalisée.')
    add_bullet(doc, 'Les horaires d\'ouverture et de fermeture du parc.')
    add_para(doc, 'Un parcours « complet » inclut tous les spectacles sélectionnés. Un parcours « partiel » '
        'n\'a pas pu intégrer tous les spectacles (manque de temps compatible).', italic=True)

    # --- Historique
    add_heading(doc, '4.1.6 Historique des Visites', 3)
    add_para(doc, 'Accessible via l\'onglet « Historique ».')
    add_bullet(doc, 'Affiche toutes les visites enregistrées avec leur date et nom.')
    add_bullet(doc, 'Appuyer sur une visite pour revoir les parcours calculés.')
    add_bullet(doc, 'Changer de parcours favori en appuyant sur l\'étoile (⭐) du parcours souhaité.')
    add_bullet(doc, 'Accéder à la carte pour le parcours favori via le bouton « Carte ».')

    # --- Carte
    add_heading(doc, '4.1.7 Vue Carte', 3)
    add_para(doc, (
        'La carte affiche le parcours favori de la visite sélectionnée sur une carte OpenStreetMap. '
        'Les marqueurs indiquent chaque étape dans l\'ordre (avec l\'heure de début). '
        'Une ligne relie les étapes pour visualiser le chemin à parcourir dans le parc.'
    ))
    add_bullet(doc, 'Nécessite une connexion internet pour charger les tuiles cartographiques.')
    add_bullet(doc, 'Requiert que la visite ait un parcours favori défini.')

    # --- Profil
    add_heading(doc, '4.1.8 Profil et Vitesse de Marche', 3)
    add_para(doc, 'Accessible via l\'onglet « Profil ».')
    add_bullet(doc, 'Affiche les informations du compte (nom, prénom, email).')
    add_bullet(doc, 'Curseur (SeekBar) pour ajuster la vitesse de marche de 2,0 à 7,0 km/h (pas de 0,5 km/h).')
    add_bullet(doc, 'Appuyer sur « Sauvegarder » pour mettre à jour la vitesse utilisée dans les calculs futurs.')
    add_bullet(doc, 'Bouton « Se déconnecter » pour quitter l\'application.')

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENTATION UTILISATEUR — GESTIONNAIRE
# ═══════════════════════════════════════════════════════════════════════════════

def build_user_manager(doc):
    add_heading(doc, '4.2 Guide Gestionnaire — Back-office Web', 2)

    add_para(doc, (
        'Ce guide s\'adresse aux gestionnaires du parc Puy du Fou qui administrent le contenu '
        'via l\'interface web back-office. URL d\'accès : '
        'http://localhost/ppe-puy-du-fou/V1/backoffice/public'
    ))

    # --- Connexion
    add_heading(doc, '4.2.1 Connexion', 3)
    add_table(doc,
        ['Étape', 'Action'],
        [
            ['1', 'Accéder à l\'URL du back-office dans un navigateur web.'],
            ['2', 'Saisir l\'adresse email et le mot de passe du compte gestionnaire.'],
            ['3', 'Cliquer sur « Se connecter ».'],
            ['4', 'En cas de succès, redirection automatique vers le tableau de bord.'],
        ],
        col_widths=[1.5, 15.5]
    )
    add_para(doc, 'Compte de test : admin@puydufou.fr / password', italic=True)

    # --- Dashboard
    add_heading(doc, '4.2.2 Tableau de Bord', 3)
    add_para(doc, (
        'Le tableau de bord présente en un coup d\'œil les compteurs principaux : '
        'nombre de spectacles, de lieux, de jours d\'ouverture et de séances configurées. '
        'Des liens rapides permettent d\'accéder directement à chaque module de gestion.'
    ))

    # --- Spectacles
    add_heading(doc, '4.2.3 Gestion des Spectacles', 3)
    add_para(doc, 'Menu : « Spectacles »')
    add_table(doc,
        ['Action', 'Comment faire'],
        [
            ['Voir la liste', 'La page affiche tous les spectacles avec leur lieu et leurs durées.'],
            ['Créer un spectacle', 'Cliquer « + Nouveau spectacle », remplir le formulaire (libellé, description, durée spectacle, durée attente, lieu associé), valider.'],
            ['Modifier un spectacle', 'Cliquer l\'icône « Éditer » sur la ligne du spectacle, modifier les champs, sauvegarder.'],
            ['Supprimer un spectacle', 'Cliquer l\'icône « Supprimer ». Confirmation demandée. Suppression en cascade (séances associées supprimées).'],
        ],
        col_widths=[4, 13]
    )

    # --- Lieux
    add_heading(doc, '4.2.4 Gestion des Lieux', 3)
    add_para(doc, 'Menu : « Lieux »')
    add_table(doc,
        ['Champ', 'Description'],
        [
            ['Nom du lieu', 'Nom affiché dans l\'application (ex: « Cinéscénie », « Restaurant Médiéval »).'],
            ['Type', 'ENUM : spectacle, zone, restaurant, hôtel, accueil, allée.'],
            ['Latitude / Longitude', 'Coordonnées GPS WGS84 pour le placement sur la carte.'],
        ],
        col_widths=[4, 13]
    )
    add_para(doc, 'Attention : la suppression d\'un lieu supprime les distances associées en cascade. '
        'Il faut reconfigurer les distances si un lieu est recréé.', italic=True)

    # --- Séances
    add_heading(doc, '4.2.5 Gestion des Séances', 3)
    add_para(doc, 'Menu : « Séances »')
    add_bullet(doc, 'Filtrer les séances par date via le sélecteur de date.')
    add_bullet(doc, 'Créer une séance : choisir le spectacle, la date, l\'heure de début et l\'heure de fin.')
    add_bullet(doc, 'Vérifier que les horaires d\'une séance sont compris dans les horaires du parc.')
    add_bullet(doc, 'Modifier ou supprimer une séance via les icônes d\'action.')

    # --- Jours
    add_heading(doc, '4.2.6 Gestion des Jours d\'Ouverture', 3)
    add_para(doc, 'Menu : « Jours d\'ouverture »')
    add_para(doc, (
        'Un « jour d\'ouverture » définit les horaires du parc pour une date donnée. '
        'Ces horaires sont utilisés par l\'algorithme de calcul pour valider qu\'une séance '
        'est accessible dans la journée.'
    ))
    add_table(doc,
        ['Champ', 'Exemple', 'Description'],
        [
            ['Date', '2026-04-11', 'Date d\'ouverture (format YYYY-MM-DD). Clé primaire — UPSERT possible.'],
            ['Heure d\'ouverture', '09:30', 'Heure à partir de laquelle les visiteurs entrent.'],
            ['Heure de fermeture', '19:30', 'Heure limite pour les spectacles.'],
        ],
        col_widths=[3.5, 3, 10.5]
    )

    # --- Distances
    add_heading(doc, '4.2.7 Gestion des Distances (Graphe)', 3)
    add_para(doc, 'Menu : « Distances »')
    add_para(doc, (
        'Le graphe des distances est la base du calcul d\'itinéraires. Chaque arête relie deux lieux '
        'avec une distance en mètres. L\'algorithme Floyd-Warshall utilise ce graphe pour calculer '
        'les plus courts chemins entre tous les lieux.'
    ))
    add_bullet(doc, 'Ajouter une arête : sélectionner lieu A, lieu B, saisir la distance en mètres, valider.')
    add_bullet(doc, 'L\'arête est bidirectionnelle (A→B et B→A créés simultanément).')
    add_bullet(doc, 'Supprimer une arête : cliquer l\'icône « Supprimer » sur la ligne de l\'arête.')
    add_bullet(doc, 'Important : si un lieu n\'est pas relié au graphe, les visiteurs ne pourront pas y être acheminés.')

    add_heading(doc, '4.2.8 Déconnexion', 3)
    add_para(doc, 'Cliquer sur « Se déconnecter » dans la barre de navigation supérieure. '
        'La session est détruite et le token JWT supprimé.')

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# INSTALLATION ET MISE EN ROUTE
# ═══════════════════════════════════════════════════════════════════════════════

def build_installation(doc):
    add_heading(doc, '5. Installation et Mise en Route', 1)

    add_heading(doc, '5.1 Prérequis', 2)
    add_table(doc,
        ['Composant', 'Version minimale', 'Rôle'],
        [
            ['Laragon / LAMP', 'Latest', 'Serveur PHP + MySQL local'],
            ['PHP', '8.0+', 'Exécution API et back-office'],
            ['MySQL', '8.0+', 'Base de données'],
            ['Android Studio', '2022+', 'Compilation de l\'application Android'],
            ['JDK', '11+', 'Compilation Java'],
        ],
        col_widths=[4, 3.5, 9.5]
    )

    add_heading(doc, '5.2 Installation de la Base de Données', 2)
    add_code(doc, '# 1. Démarrer MySQL (Laragon)')
    add_code(doc, '# 2. Ouvrir phpMyAdmin ou HeidiSQL')
    add_code(doc, '# 3. Créer la base de données : ppe_puy_du_fou')
    add_code(doc, '# 4. Importer le fichier SQL :')
    add_code(doc, 'mysql -u root ppe_puy_du_fou < V1/db/ppe_puy_du_fou.sql')
    doc.add_paragraph()

    add_heading(doc, '5.3 Configuration de l\'API', 2)
    add_para(doc, 'Fichier : api/config/config.php')
    add_code(doc, "define('DB_HOST',       '127.0.0.1:3306');")
    add_code(doc, "define('DB_NAME',       'ppe_puy_du_fou');")
    add_code(doc, "define('DB_USER',       'root');")
    add_code(doc, "define('DB_PASS',       '');")
    add_code(doc, "define('JWT_SECRET',    'change-me-in-production-please');")
    add_code(doc, "define('JWT_ALGO',      'HS256');")
    add_code(doc, "define('JWT_EXPIRES',   86400);")
    doc.add_paragraph()
    add_para(doc, 'URL d\'accès API : http://localhost/ppe-puy-du-fou/V1/api/public')

    add_heading(doc, '5.4 Configuration de l\'Application Android', 2)
    add_para(doc, 'Fichier : android/app/src/main/java/.../api/ApiClient.java')
    add_code(doc, 'private static final String BASE_URL = "http://10.0.2.2/ppe-puy-du-fou/V1/api/public";')
    add_code(doc, '// 10.0.2.2 = adresse de la machine hôte depuis l\'émulateur Android')
    add_code(doc, '// Sur appareil physique : remplacer par l\'IP locale du PC (ex: 192.168.1.x)')
    doc.add_paragraph()

    add_heading(doc, '5.5 Comptes de Test', 2)
    add_table(doc,
        ['Rôle', 'Email', 'Mot de passe', 'Usage'],
        [
            ['Gestionnaire', 'admin@puydufou.fr', 'password', 'Back-office + API admin'],
            ['Visiteur 1', 'jean.dupont@email.fr', 'password', 'Application Android'],
            ['Visiteur 2', 'marie.curie@email.fr', 'password', 'Application Android'],
        ],
        col_widths=[3, 5, 3, 6]
    )

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# UPLOAD GOOGLE DRIVE (OPTIONNEL)
# ═══════════════════════════════════════════════════════════════════════════════

def upload_to_google_drive(docx_path):
    """
    Upload le DOCX vers Google Drive et le convertit en Google Doc.
    Nécessite : credentials.json dans le même dossier (OAuth2 Desktop app).
    """
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    import pickle

    SCOPES     = ['https://www.googleapis.com/auth/drive.file']
    CREDS_FILE = os.path.join(SCRIPT_DIR, 'credentials.json')
    TOKEN_FILE  = os.path.join(SCRIPT_DIR, 'token.pickle')

    if not os.path.exists(CREDS_FILE):
        print('\n[ERREUR] credentials.json introuvable.')
        print('Télécharger depuis : Google Cloud Console → APIs & Services → Credentials')
        print('Créer un OAuth 2.0 Client ID (Desktop App), télécharger JSON → renommer credentials.json')
        return None

    creds = None
    if os.path.exists(TOKEN_FILE):
        with open(TOKEN_FILE, 'rb') as f:
            creds = pickle.load(f)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow  = InstalledAppFlow.from_client_secrets_file(CREDS_FILE, SCOPES)
            creds = flow.run_local_server(port=0)
        with open(TOKEN_FILE, 'wb') as f:
            pickle.dump(creds, f)

    service = build('drive', 'v3', credentials=creds)

    file_metadata = {
        'name':     'PPE Puy du Fou — Documentation Complète',
        'mimeType': 'application/vnd.google-apps.document',
    }
    media = MediaFileUpload(
        docx_path,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        resumable=True
    )

    print('\nUpload vers Google Drive en cours...')
    file = service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id,webViewLink'
    ).execute()

    url = file.get('webViewLink')
    print(f'[OK] Google Doc créé : {url}')
    return url


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    upload = '--upload' in sys.argv

    print('Génération de la documentation PPE Puy du Fou V1...')

    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Construire toutes les sections
    build_cover(doc)
    build_presentation(doc)
    build_mcd(doc)
    build_api(doc)
    build_backoffice(doc)
    build_android(doc)
    build_user_visitor(doc)
    build_user_manager(doc)
    build_installation(doc)

    # Sauvegarder
    doc.save(OUTPUT_FILE)
    print(f'[OK] DOCX généré : {OUTPUT_FILE}')

    # Upload optionnel
    if upload:
        url = upload_to_google_drive(OUTPUT_FILE)
        if url:
            print(f'\nGoogle Doc disponible : {url}')
    else:
        print('\nPour uploader vers Google Drive :')
        print('  1. Placer credentials.json dans ce dossier (OAuth2 Desktop)')
        print('  2. Relancer : python create_documentation.py --upload')
        print('\nOu importer manuellement le DOCX dans Google Drive (Fichier > Importer).')


if __name__ == '__main__':
    main()
